from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


PROJECT_ROOT = Path(r"F:\通信达量化\tdx_modular_quant_project_v2_all_instruments")
OUTPUT = PROJECT_ROOT / "reports" / "量化选股_持仓数量_卖出公式说明书.docx"
SKILL_ROOT = Path(r"C:\Users\Ziyi Wang\.codex\plugins\cache\openai-primary-runtime\documents\26.802.11031\skills\documents")
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203040"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D0D5DD", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=11, bold=False, color=INK, italic=False, math=False):
    run.font.name = "Cambria Math" if math else "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Cambria Math" if math else "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Cambria Math" if math else "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill=LIGHT_GRAY):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_formula(doc, formula, note=None):
    p = doc.add_paragraph(style="Formula")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(p, LIGHT_GRAY)
    run = p.add_run(formula)
    set_run_font(run, size=11.5, bold=True, color=DARK_BLUE, math=True)
    if note:
        n = doc.add_paragraph(style="Formula Note")
        r = n.add_run(note)
        set_run_font(r, size=9.5, color=MUTED)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_note(doc, label, text, color=GOLD):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF8E8")
    set_cell_border(cell, "E3C66A", "8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(label + "：")
    set_run_font(lead, size=10, bold=True, color=color)
    body = p.add_run(text)
    set_run_font(body, size=10, color=INK)
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=160,
                         cell_margins_dxa={"top": 110, "bottom": 110, "start": 160, "end": 160})
    return table


def add_param_table(doc, rows, widths=(1500, 2200, 1300, 4360)):
    table = doc.add_table(rows=1, cols=4)
    headers = ("参数/符号", "代码字段", "当前值", "含义与设置方法")
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = str(text)
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            for run in p.runs:
                set_run_font(run, size=9.0, color=INK)
    apply_table_geometry(table, list(widths), table_width_dxa=9360, indent_dxa=120,
                         cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120})
    return table


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    formula = doc.styles.add_style("Formula", 1)
    formula.font.name = "Cambria Math"
    formula.font.size = Pt(11.5)
    formula.paragraph_format.space_before = Pt(5)
    formula.paragraph_format.space_after = Pt(3)
    formula.paragraph_format.left_indent = Inches(0.15)
    formula.paragraph_format.right_indent = Inches(0.15)
    formula.paragraph_format.line_spacing = 1.15
    formula.paragraph_format.keep_together = True

    formula_note = doc.styles.add_style("Formula Note", 1)
    formula_note.font.name = "Calibri"
    formula_note.font.size = Pt(9.5)
    formula_note.font.italic = True
    formula_note.font.color.rgb = RGBColor.from_string(MUTED)
    formula_note.paragraph_format.space_before = Pt(0)
    formula_note.paragraph_format.space_after = Pt(6)
    formula_note.paragraph_format.left_indent = Inches(0.18)
    formula_note.paragraph_format.right_indent = Inches(0.18)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("量化策略计算参考 · 选股 / 持仓数量 / 卖出")
    set_run_font(hr, size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("第 ")
    set_run_font(fr, size=9, color=MUTED)
    add_field(fp, "PAGE")
    fr2 = fp.add_run(" 页")
    set_run_font(fr2, size=9, color=MUTED)


def build_document():
    doc = Document()
    configure_document(doc)

    # Cover / opening block: editorial-cover archetype, compacted for a reference manual.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("量化策略公式说明书")
    set_run_font(r, size=28, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("选股计算方式 · 动态持仓数量 · 卖出计算方式")
    set_run_font(r, size=15, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("适用口径：small_capital_lean / mainline V3 / E4")
    set_run_font(r, size=10.5, color=MUTED)

    add_note(doc, "阅读提示", "本说明书解释当前系统的实际计算逻辑。评分不是收益率，参数建议是受控研究建议，不代表可以直接调参或实盘上线。")
    doc.add_paragraph()
    add_param_table(doc, [
        ("策略口径", "capital_profile", "Lean 2万元", "小资金、允许现金、动态持仓数量。"),
        ("选股版本", "strategy_logic_version", "mainline V3", "74因子柜评分后，再按人民币净效用和整数手数选择。"),
        ("卖出阶段", "scap_exit_stage", "E4", "累计开放亏损、信号、陈旧和利润保护退出。"),
        ("正常买入频率", "portfolio_normal_rebalance_frequency", "monthly", "正常新买入按月度计划；风险卖出仍每日检查。"),
    ], widths=(1500, 2500, 1400, 3960))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("版本日期：2026-08-03")
    set_run_font(r, size=9.5, color=MUTED)

    doc.add_page_break()

    doc.add_heading("1 选股计算方式", level=1)
    add_body(doc, "选股分为四层：因子标准化、经济族与角色合成、候选人民币效用、整数组合优化。排名靠前只是获得候选资格，不能直接等同于买入。")

    doc.add_heading("1.1 因子横截面标准化", level=2)
    add_formula(doc, "FactorScore(i,f,t) = PctRank_t[ RawFactor(i,f,t) ]",
                "同一决策日内，把第 f 个因子对股票 i 的原始值转换为 0～1 百分位。越接近 1，表示该股票在这一因子上的当日排名越高。")
    add_param_table(doc, [
        ("i", "symbol", "股票代码", "被评价的股票。"),
        ("f", "model_name", "74个因子之一", "当前冻结因子柜内的因子模型。因子数量不能临时增减，否则结果不可直接比较。"),
        ("t", "date", "决策日", "只使用截至该日可见的数据，禁止使用未来数据。"),
        ("RawFactor", "predicted_return_5d", "因子原始值", "因子模型输出；不同因子量纲可能不同。"),
        ("PctRank", "factor_score", "0～1", "当日横截面百分位。用于统一不同因子的量纲。"),
    ])

    doc.add_heading("1.2 相近因子和经济族压缩", level=2)
    add_formula(doc, "RelativeScore(i,c) = Median_{f in c}[ FactorScore(i,f) ]")
    add_formula(doc, "FamilyScore(i,g) = Median_{c in g}[ RelativeScore(i,c) ]",
                "先对相近因子取中位数，再对同一经济逻辑下的聚类取中位数，防止多个相似因子重复投票。")
    add_param_table(doc, [
        ("c", "empirical_cluster_key", "因子聚类", "语义近亲因子自动归组；额外经验聚类要求绝对 Spearman 相关不低于 0.90。"),
        ("g", "economic_family", "经济因子族", "例如动量、反转、规模、流动性、风险等经济含义。"),
        ("相关阈值", "empirical_cluster_threshold", "0.90", "建议保持冻结。降低会合并更多因子；提高会让相似因子获得更多重复权重。"),
        ("最少样本", "empirical_cluster_min_observations", "30", "低于30个有效观察时，不使用经验相关聚类。"),
        ("Median", "median", "中位数", "比平均值更不容易被单个极端因子支配。"),
    ])

    doc.add_heading("1.3 角色评分", level=2)
    add_formula(doc, "FamilyWeight(i,g) = min[ 1 / ActiveFamilyCount(i,r), 0.25 ]")
    add_formula(doc, "Coverage(i,r) = clip[ ActiveFamilyCount / ConfiguredFamilyCount, 0, 1 ]")
    add_formula(doc, "RoleScore(i,r) = clip{ 0.5 + Coverage × Σ_g[ FamilyWeight × (FamilyScore - 0.5) ], 0, 1 }",
                "0.5 是无证据时的中性分。覆盖不足会把角色评分收缩回 0.5；单一经济族最多取得 25% 权重。")
    add_param_table(doc, [
        ("r", "primary_role", "决策角色", "严格入场、代理入场、择时、风险、流动性、持有验证或卖出安全。"),
        ("0.5", "neutral_score", "中性值", "高于0.5为支持，低于0.5为不支持。"),
        ("单族上限", "max_family_share", "0.25", "单一经济族最多贡献25%。建议冻结；提高会增加单一风格支配风险。"),
        ("Coverage", "*_coverage", "0～1", "角色实际覆盖率。严格入场覆盖率为0时，不具有正式新买入资格。"),
        ("RoleScore", "cabinet_*_score", "0～1", "各角色的最终连续评分，不是未来收益率。"),
    ])

    doc.add_heading("1.4 最终柜评分", level=2)
    add_formula(doc, "StrictAuthority = StrictCoverage × StrictWeight")
    add_formula(doc, "ProxyAuthority = (1 - StrictCoverage) × ProxyCoverage")
    add_formula(doc, "BaseEntry = (StrictScore×StrictAuthority + ProxyScore×ProxyAuthority) / (StrictAuthority+ProxyAuthority)")
    add_formula(doc, "TimingAdjustment = clip[ TimingWeight × (TimingScore - 0.5), -0.15, 0.15 ]")
    add_formula(doc, "LiquidityPenalty = clip[ 0.20 × max(0, 0.5-LiquidityScore), 0, 0.10 ]")
    add_formula(doc, "CabinetFinalScore = clip[ BaseEntry + TimingAdjustment - LiquidityPenalty, 0, 1 ]",
                "当前 StrictWeight=1、TimingWeight=0，因此代理因子只填补严格因子缺失，择时分不直接改变最终排序。")
    add_param_table(doc, [
        ("StrictWeight", "strict_weight", "1.0", "严格入场证据权重。建议保持1.0；调低会无依据地放大代理证据。"),
        ("TimingWeight", "timing_weight", "0.0", "当前仅审计择时，不让择时直接改排名。若要开启，必须做PIT和样本外验证。"),
        ("±0.15", "timing_adjustment_cap", "±0.15", "择时最大加减分。即使未来启用，也不能无限支配主评分。"),
        ("0.20", "liquidity_penalty_slope", "0.20", "流动性低于0.5后的扣分斜率。"),
        ("0.10", "liquidity_penalty_cap", "0.10", "流动性最大扣分。"),
        ("最终分", "cabinet_native_final_score", "0～1", "只用于相对排序；不能解释为上涨概率或收益率。"),
    ])

    doc.add_heading("1.5 一手可买性和人民币效用", level=2)
    add_formula(doc, "OneLotNotional = Price × MinimumBuyShares")
    add_formula(doc, "OneLotWeight = OneLotCashRequired / NAV")
    add_formula(doc, "CandidateUtility = Notional × DecisionExpectedReturn - TotalLifecycleCost - RiskPenaltyAmount",
                "只有覆盖有效、未处于退出/冷却状态、一手现金可负担、单股不穿透硬上限，并且人民币效用为正的股票，才进入整数优化。")
    add_param_table(doc, [
        ("Price", "close_nominal / close", "元/股", "决策时可见的名义价格。正式执行还要使用下一交易日可成交价格。"),
        ("MinimumBuyShares", "minimum_buy_quantity", "通常100股", "按股票板块和交易日期的真实交易规则决定。"),
        ("NAV", "nominal_nav", "人民币元", "现金加全部持仓市值。"),
        ("DecisionExpectedReturn", "comparable_alpha_lcb", "收益比例", "使用校准后的保守收益下界，不使用0～1柜评分代替收益率。"),
        ("预测期限", "comparable_value_horizon_days", "通常10日", "收益预测对应的交易日长度；不能把不同期限的预测直接相加。"),
        ("现金缓冲", "min_cash_buffer", "1,000元", "Lean档至少保留的现金。调高更保守，调低增加满仓和费用风险。"),
    ])

    doc.add_heading("1.6 经济订单门", level=2)
    add_formula(doc, "ConservativeGrossProfit = ExpectedReturnLCB × Notional")
    add_formula(doc, "RobustNetProfit = ConservativeGrossProfit - LifecycleCost")
    add_formula(doc, "CostShare = LifecycleCost / ConservativeGrossProfit")
    add_formula(doc, "EconomicOrderPass = (CostShare ≤ 0.30) AND (RobustNetProfit ≥ 15元)",
                "完整生命周期费用必须被保守毛利润覆盖，且保守净利润至少15元。")
    add_param_table(doc, [
        ("ExpectedReturnLCB", "comparable_alpha_lcb", "保守下界", "必须来自有PIT权威的校准预测；缺失或未授权时不允许产生正效用。"),
        ("LifecycleCost", "total_lifecycle_cost_amount", "人民币元", "买入费用＋未来卖出费用；当前预期补仓/替换概率均为0。"),
        ("费用占比上限", "scap_max_lifecycle_cost_to_gross_profit_ratio", "0.30", "费用最多吃掉30%的保守毛利润。建议通过滚动样本外费用敏感性研究设置。"),
        ("最低稳健利润", "scap_minimum_robust_profit_hurdle_amount", "15元", "防止理论正收益但人民币利润太小。应结合账户规模、最低佣金和成交频率设置。"),
        ("候选最低佣金", "scap_candidate_minimum_commission", "5元/边", "经济性规划假设。应改为实际券商费率，不能随意取低。"),
    ])

    doc.add_heading("1.7 最终选中公式", level=2)
    add_formula(doc, "PortfolioObjective = IncrementalRobustWealth - ScenarioRiskPenalty - ThesisPenalty - ConcentrationPenalty")
    add_formula(doc, "LexicographicKey = (RobustProfit, -DeploymentGap, Breadth, ExpectedProfit, -Downside, -Cost)",
                "系统先最大化成本后稳健人民币利润；只有第一目标相同时，才依次比较部署缺口、分散度、普通预期利润、下行风险和费用。")
    add_note(doc, "参数设置原则", "选股评分参数、收益校准参数和费用参数必须分别管理。0～1评分不能直接乘本金当作收益；任何权重或阈值变化都应冻结同一数据窗、因子柜、成本和PIT状态做受控样本外比较。")

    doc.add_page_break()
    doc.add_heading("2 持仓数量计算方式", level=1)
    add_body(doc, "当前 Lean 档采用动态持仓数量 K。K 表示当日最多允许存在多少个股票名称，不等于系统必须把槽位全部填满。")

    doc.add_heading("2.1 可投资资金", level=2)
    add_formula(doc, "RiskBudget = NAV × RiskExposureCeiling")
    add_formula(doc, "CurrentInvested = NAV × CurrentExposure")
    add_formula(doc, "RiskRoom = max[ RiskBudget - CurrentInvested, 0 ]")
    add_formula(doc, "SpendableCash = max{ min[ Cash-MinCashBuffer, RiskRoom ], 0 }",
                "可用于新增仓位的现金，同时受真实现金和剩余风险额度限制。")
    add_param_table(doc, [
        ("NAV", "nominal_nav", "人民币元", "账户现金＋全部持仓市值。"),
        ("Cash", "cash_amount", "人民币元", "当前可用现金，不含尚未完成卖出的预计回款。"),
        ("RiskExposureCeiling", "risk_exposure_ceiling", "0～1", "当日风险模块授权的总仓位上限。不能由选股模块自行提高。"),
        ("CurrentExposure", "current_exposure", "0～1", "当前持仓市值/NAV。"),
        ("MinCashBuffer", "min_cash_buffer", "1,000元", "最低现金缓冲。建议按最低佣金、计划成交数和应急资金设置。"),
    ])

    doc.add_heading("2.2 整手、费用和候选容量", level=2)
    add_formula(doc, "K_lot = HeldNames + max k, subject to Σ_{j=1..k} OneLotCash(j) ≤ SpendableCash")
    add_formula(doc, "K_economic = HeldNames + max k, subject to Σ EconomicRequiredCash(j) ≤ SpendableCash")
    add_formula(doc, "K_candidate = HeldNames + EligibleNewNameCount",
                "K_lot 只检查整手现金；K_economic 进一步要求每个候选通过完整费用和最低稳健利润门。")
    add_param_table(doc, [
        ("HeldNames", "current_symbols", "实际持仓数", "当前账户中股数大于0的股票数量。"),
        ("OneLotCash", "mainline_v3_one_lot_cash_required", "人民币元", "一手名义金额加买入侧费用。"),
        ("EconomicRequiredCash", "economic_required_cash", "人民币元", "该候选首次出现正稳健利润且通过费用门的最小整数手数所需现金。"),
        ("EligibleNewNameCount", "eligible_symbol_count", "整数", "排除已有持仓后，仍有有效一手金额的候选数量。"),
    ])

    doc.add_heading("2.3 动态 K 总公式", level=2)
    add_formula(doc, "K_raw = min[ K_economic, K_candidate, K_search, K_user ]")
    add_formula(doc, "K_effective = max[ K_raw, HeldNames, 1 ]",
                "K_user 未设置时从最小值公式中删除。已有持仓超过新上限时，旧持仓可以祖父化保留，但不得继续扩大超限。")
    add_param_table(doc, [
        ("K_search", "scap_search_position_cap", "32", "计算搜索上限，不是金融目标。提高会增加计算量，不能用来强迫多持股。"),
        ("K_user", "user_hard_position_cap", "当前为空", "用户可选硬上限。若设置，必须为正整数；不建议在没有容量研究时随意固定为5。"),
        ("K_effective", "effective_position_cap", "每日动态", "真实有效持仓名称上限。取现金、费用、风险、候选和搜索能力的共同最小值。"),
        ("祖父化超限", "grandfathered_excess_names", "动态", "当上限下降但旧仓仍存在时的超限名称数；必须披露并禁止新增恶化。"),
    ])

    doc.add_heading("2.4 单股仓位配套上限", level=2)
    add_formula(doc, "EqualWeight = TargetExposure / K_sizing")
    add_formula(doc, "SoftCap = min[ 25%, 1.5 × EqualWeight ]")
    add_formula(doc, "HardCap = max{ SoftCap, min[ 40%, 2.3 × EqualWeight ] }")
    add_formula(doc, "CurrentWeight(i) + NewOrderWeight(i) ≤ HardCap",
                "软上限以上会产生集中度处罚；硬上限以上直接拒绝。即使只买一手，也不能穿透硬上限。")
    add_param_table(doc, [
        ("TargetExposure", "target_exposure", "0～1", "计划总仓位，不是风险授权上限；不得混淆。"),
        ("K_sizing", "sizing_reference_positions", "每日动态", "用于分散计算的参考持股数，不等于必须持有数量。"),
        ("绝对软上限", "scap_single_position_soft_cap", "25%", "达到后开始处罚。应结合账户规模和可买一手金额设置。"),
        ("绝对硬上限", "retail_single_position_cap", "40%", "不可穿透的单股结构上限。不要为追求收益放宽。"),
        ("软倍数", "soft_equal_weight_multiple", "1.5", "允许单股达到等权仓位的1.5倍。"),
        ("硬倍数", "hard_equal_weight_multiple", "2.3", "允许单股达到等权仓位的2.3倍，但仍受40%绝对上限约束。"),
    ])
    add_note(doc, "设置建议", "动态 K 应由真实现金、一手金额、完整生命周期费用、风险额度和有效候选共同决定。搜索上限、用户硬上限和实际持股数必须分开披露，不能把“最多可搜索32只”解释为“应该持有32只”。")

    doc.add_page_break()
    doc.add_heading("3 卖出计算方式", level=1)
    add_body(doc, "卖出不是一条固定止损线，而是多个退出家族每天计算后，由唯一退出仲裁器选择一个权威原因。当前 Lean 使用 E4，亏损退出需要2日确认，灾难止损立即生效；买后失败仅作纸面诊断。")

    doc.add_heading("3.1 持仓基础变量", level=2)
    add_formula(doc, "EntryPrice_new = (EntryPrice_old×OldShares + BuyPrice×NewShares) / (OldShares+NewShares)")
    add_formula(doc, "Unrealized = CurrentPrice / EntryPrice - 1")
    add_formula(doc, "MFE = HighestPriceSinceEntry / EntryPrice - 1")
    add_formula(doc, "MAE = LowestPriceSinceEntry / EntryPrice - 1")
    add_formula(doc, "Giveback = (MFE-Unrealized) / MFE, when MFE>0; otherwise 0",
                "MFE 是持仓期间最高浮盈，MAE 是最深浮亏，Giveback 是从最高利润回吐的比例。")
    add_param_table(doc, [
        ("EntryPrice", "position_entry_price", "元/股", "首次买入价；补仓后按股数加权更新。"),
        ("CurrentPrice", "current_price", "元/股", "决策日可见价格。"),
        ("HighestPriceSinceEntry", "peak_price", "元/股", "入场以来观测到的最高价格。"),
        ("LowestPriceSinceEntry", "trough_price", "元/股", "入场以来观测到的最低价格。"),
        ("MFE/MAE", "position_mfe / position_mae", "收益比例", "分别衡量最好和最差持仓路径。"),
        ("Giveback", "position_giveback_from_peak", "0～1以上", "峰值利润回吐比例；越高表示利润保护压力越大。"),
    ])

    doc.add_heading("3.2 卖出触发的成本后收益", level=2)
    add_formula(doc, "TriggerCostRate = 2×Commission + 2×Slippage + StampDuty + 2×TransferFee")
    add_formula(doc, "TriggerCostRate = 2×0.03% + 2×0.05% + 0.05% + 2×0.001% = 0.212%")
    add_formula(doc, "NetUnrealized = Unrealized - 0.212%;   NetMFE = MFE - 0.212%",
                "该比例用于生命周期触发；真实下单仍按逐笔最低佣金、滑点、印花税、过户费和市场冲击计算。")
    add_param_table(doc, [
        ("佣金率", "COMMISSION_RATE", "0.03%/边", "应按券商真实费率设置。"),
        ("滑点率", "SLIPPAGE_RATE", "0.05%/边", "应根据成交偏差和流动性压力测试设置。"),
        ("印花税", "STAMP_DUTY_RATE", "0.05%/卖出", "仅卖出侧，并按历史交易日期费率处理。"),
        ("过户费", "TRANSFER_FEE_RATE", "0.001%/边", "买卖两边计入。"),
    ])

    doc.add_heading("3.3 灾难止损和自适应软止损", level=2)
    add_formula(doc, "DisasterExit = Held AND HoldingDays≥1 AND NetUnrealized≤-18%")
    add_formula(doc, "AdaptiveSoftStop = clip[ -16% + 4%×TailRiskProxy, -16%, -12% ]")
    add_formula(doc, "AdaptiveLossExit = Held AND HoldingDays≥3 AND NetUnrealized≤AdaptiveSoftStop for 2 consecutive decision days")
    add_formula(doc, "LossExit = DisasterExit OR AdaptiveLossExit",
                "灾难止损立即退出；普通自适应亏损必须至少持有3日并连续2日确认。尾部风险越高，软止损从-16%逐步收紧到-12%。")
    add_param_table(doc, [
        ("灾难止损", "scap_loss_disaster_stop", "-18%", "单日安全断路器。应基于极端损失承受能力设置，不建议放宽。"),
        ("软止损基础", "scap_loss_soft_base", "-16%", "尾部风险最低时的普通止损线。"),
        ("尾部收紧", "scap_loss_tail_tightening", "4%", "TailRiskProxy 从0升至1时最多收紧4个百分点。"),
        ("最紧软止损", "scap_loss_soft_tightest", "-12%", "自适应止损不能比-12%更紧。"),
        ("确认日数", "scap_loss_stop_confirmation_days", "2日", "过滤单日噪声；增大更迟钝，减小更敏感。"),
        ("TailRiskProxy", "tail_risk_proxy", "0～1", "尾部风险代理分数。只能使用决策日可见信息。"),
    ])

    doc.add_heading("3.4 硬止盈回撤", level=2)
    add_formula(doc, "ProfitHardStopExit = Held AND NetMFE≥12% AND NetUnrealized≥3% AND (NetMFE-NetUnrealized)/NetMFE≥55%",
                "曾经取得至少12%的成本后最大浮盈、目前仍保留至少3%净利润，但已回吐峰值利润55%以上时整仓退出。")
    add_param_table(doc, [
        ("启动利润", "scap_profit_protection_arm", "12%", "盈利保护开始生效的最低 NetMFE。"),
        ("最低剩余利润", "scap_profit_protection_min_net_profit", "3%", "触发硬止盈时至少仍保留的净利润。"),
        ("允许回吐", "scap_profit_protection_giveback", "55%", "峰值利润最多允许回吐的比例。调低更早止盈，调高更容忍趋势波动。"),
    ])

    doc.add_heading("3.5 动态利润回吐", level=2)
    add_formula(doc, "BaseGiveback = 40% if MFE<15%; 30% if 15%≤MFE<25%; 25% if MFE≥25%")
    add_formula(doc, "OrderflowDecay = max[0.55-OrderflowScore, 0] / 0.55")
    add_formula(doc, "DynamicGiveback = max[ BaseGiveback - 0.30×max(0.55-TrendDirection,0) - 0.12×PeakDecay - 0.08×OrderflowDecay, 8% ]")
    add_formula(doc, "ProtectionPressure = clip[ 0.40×Giveback/DynamicGiveback + 0.25×PeakDecay + 0.20×(1-TrendDirection) + 0.15×OrderflowDecay, 0, 1 ]")
    add_formula(doc, "ProfitGivebackExit = Held AND MFE≥8% AND Giveback≥DynamicGiveback AND ProtectionPressure≥0.70",
                "盈利越高、趋势越弱、峰值衰退越严重，系统允许的回吐越少。")
    add_param_table(doc, [
        ("基础回吐", "GOVERNANCE_PROFIT_GIVEBACK_1/2/3", "40%/30%/25%", "分别对应低于15%、15%～25%、不低于25%的MFE区间。"),
        ("最低回吐线", "dynamic_giveback_floor", "8%", "动态公式再弱也不能低于8%，防止过度敏感。"),
        ("利润保护起点", "GOVERNANCE_PROFIT_PROTECT_TRIGGER_1", "8%", "MFE低于8%时不执行普通利润回吐保护。"),
        ("压力阈值", "profit_protection_pressure_threshold", "0.70", "综合回吐、趋势、峰值和订单流后的退出阈值。"),
        ("TrendDirection", "trend_direction_score", "0～1", "趋势健康度，越高越健康。"),
        ("PeakDecay", "peak_decay_score", "0～1", "价格从峰值衰退的程度，越高越危险。"),
        ("OrderflowScore", "orderflow_candidate_score", "0～1", "订单流支持分，低于0.55开始产生衰减。"),
    ])

    doc.add_heading("3.6 信号、论点和陈旧退出", level=2)
    add_formula(doc, "RawSignalFailure = Held AND HoldingDays≥GraceDays AND EntryScore<0.45 AND TrendScore<0.45")
    add_formula(doc, "ThesisFailure = Held AND HoldingDays≥GraceDays AND EntrySupport≥0.45 AND CurrentSupport<0.35 AND SupportDecay≥0.20")
    add_formula(doc, "DowntrendExit = Held AND HoldingDays≥GraceDays AND DowntrendScore≥0.75 AND FollowThroughScore<0.45")
    add_formula(doc, "StaleContext = Held AND MFE≤3% AND AlphaQualityDrop≥10% AND LiquidityDecay≥20%")
    add_formula(doc, "StaleExit = StaleContext AND HoldingDays≥30 AND Unrealized≤0",
                "价值、成长、现金流质量、盈利质量论点的 GraceDays=20；其他论点 GraceDays=10。当前 Lean 的信号失效确认日数为1日。")
    add_param_table(doc, [
        ("GraceDays", "thesis_grace_days", "10或20日", "质量/价值类逻辑给20日，其余10日。建议按策略预期持有周期设定。"),
        ("入场衰减线", "GOVERNANCE_ENTRY_MATRIX_EXIT_DECAY_THRESHOLD", "0.45", "当前入场支持低于0.45视为明显衰减。"),
        ("下跌趋势线", "GOVERNANCE_DOWNTREND_DECAY_EXIT", "0.75", "下跌衰减分达到0.75触发强风险信号。"),
        ("跟随确认线", "follow_through_threshold", "0.45", "低于0.45说明后续价格/成交确认不足。"),
        ("陈旧MFE上限", "GOVERNANCE_STALE_EXIT_MAX_MFE", "3%", "持仓期间最高只赚到3%以内。"),
        ("Alpha下降", "GOVERNANCE_STALE_EXIT_MIN_ALPHA_DROP", "10%", "当前Alpha质量相对买入时下降至少0.10。"),
        ("流动性下降", "GOVERNANCE_STALE_EXIT_MIN_LIQUIDITY_DECAY", "20%", "成交健康显著恶化。"),
        ("陈旧退出日数", "GOVERNANCE_STALE_EXIT_DAYS", "30日", "满足陈旧背景且持有至少30日、当前不盈利时退出。"),
    ])

    doc.add_heading("3.7 唯一卖出仲裁与执行边界", level=2)
    add_formula(doc, "ExitState = 1 if AuthorizedExitReason is not empty; otherwise 0")
    add_formula(doc, "TargetShares = 0 for full-liquidation exit reasons",
                "多个卖出条件同时命中时，只选择一个规范卖出原因。卖出优先于补仓和新买入，但真实成交仍必须通过T+1、停牌、跌停、可卖库存和订单规则。")
    add_param_table(doc, [
        ("E4阶段", "scap_exit_stage", "E4", "累计开放信号、陈旧、买后失败家族、亏损和利润保护授权；Lean V3仍对买后失败执行纸面降级。"),
        ("整仓原因", "FULL_LIQUIDATION_REASONS", "目标股数0", "亏损、硬止盈、信号失效、论点失效、陈旧等规范原因。"),
        ("T+1", "sellable_shares", "生产硬约束", "当日买入的股票不能当日卖出，卖出数量不得超过可卖库存。"),
        ("买后失败", "post_entry_failure_exit", "Lean纸面诊断", "当前 small_capital_lean + mainline V3 不据此真实卖出，避免旧软评分重新取得交易权。"),
    ])

    doc.add_heading("4 参数如何设置", level=1)
    add_body(doc, "参数不能根据一次回测结果逐个追涨杀跌式修改。建议按以下顺序设置，并始终保留未调参的样本外窗口。")
    add_param_table(doc, [
        ("第一层", "交易硬参数", "真实值", "佣金、最低佣金、印花税、过户费、一手数量、T+1必须来自券商和交易所规则。"),
        ("第二层", "账户约束", "按资金设置", "现金缓冲、绝对单股上限、用户持仓硬上限由资金规模和风险承受能力决定。"),
        ("第三层", "容量参数", "动态计算", "K应由现金、一手、完整费用、风险和候选共同决定，不应先拍脑袋固定。"),
        ("第四层", "收益与费用门", "滚动样本外", "最低稳健利润和费用占比上限要用真实成交成本、不同价格股票和滚动窗口校准。"),
        ("第五层", "卖出阈值", "预登记实验", "止损、止盈和陈旧期必须按同一数据、代码、成本、PIT状态做单变量或正交实验。"),
        ("禁止做法", "同窗反复调参", "禁止", "不能在同一个回测窗口反复修改阈值，再把最好结果当作未来有效证据。"),
    ], widths=(1350, 2050, 1400, 4560))
    add_note(doc, "客观建议", "当前参数只能视为工程上已实现、研究上待验证。正式修改前应先冻结目标函数和比较口径，再进行滚动样本外、费用压力、极端行情和独立复核；20日工程验收不能证明参数最优或可以实盘上线。", color=RED)

    doc.add_heading("5 口径边界", level=1)
    add_body(doc, "本说明书只覆盖选股计算、动态持仓数量和卖出计算。它不展开74个原始因子各自的生成公式，也不替代因子柜清单、执行会计、公司行动、税务、基准和正式上线验收文档。若因子柜、代码版本、费用或PIT状态发生变化，本文中的运行口径也必须同步更新。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "量化选股、持仓数量与卖出公式说明书"
    doc.core_properties.subject = "small_capital_lean / mainline V3 / E4"
    doc.core_properties.author = "量化研究项目"
    doc.core_properties.keywords = "选股, 动态K, 持仓上限, 卖出公式, 参数说明"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
